import json
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ArtisanBlueprint Brand Colors
BRAND_COPPER = RGBColor(198, 138, 107)  # #C68A6B
BRAND_DARK = RGBColor(26, 17, 16)       # #1A1110

def insert_hr(paragraph, color='1A1110'):
    """Inserts a horizontal rule (border) under a paragraph."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_cell_background(cell, fill):
    """Set background color for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:shd')
    tcBorders.set(qn('w:fill'), fill)
    tcPr.append(tcBorders)

def add_page_border(section, color='C68A6B'):
    """Add a page border to a section."""
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')  # 1pt size
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), color)
        pgBorders.append(border)
    
    sectPr.append(pgBorders)

from json_repair import repair_and_parse_json

def generate_premium_pdf(plan_json_str, page_to_images=None, docx_images_dict=None, output_filename="Premium_Plan.docx", custom_img_dir=None):
    """
    Generates an elegantly styled DOCX file aligned with ArtisanBlueprint branding.
    """
    if isinstance(plan_json_str, str):
        plan_data = repair_and_parse_json(plan_json_str)
    else:
        plan_data = plan_json_str
        
    project_name = plan_data.get("project_title", plan_data.get("project_name", "Woodworking Plan"))

    doc = Document()
    
    # --- ADD BRANDED FOOTER AND BORDER ---
    section = doc.sections[0]
    add_page_border(section, color='C68A6B')
    
    footer = section.footer
    footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_p.text = ""
    
    # Branded text
    run_brand = footer_p.add_run("ArtisanBlueprint  |  ")
    run_brand.font.size = Pt(10)
    run_brand.font.color.rgb = BRAND_COPPER
    run_brand.bold = True
    
    run1 = footer_p.add_run(f"{project_name} - Page ")
    run1.font.size = Pt(10)
    run1.font.color.rgb = RGBColor(120, 120, 120)
    
    run2 = footer_p.add_run()
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(120, 120, 120)
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    run2._r.append(fldChar3)
    
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_heading(text, level=2, center=False):
        p = doc.add_paragraph()
        run = p.add_run(text.upper() if level > 1 else text)
        run.bold = True
        
        if level == 1:
            # Hero Title: Copper
            run.font.size = Pt(24)
            run.font.color.rgb = BRAND_COPPER
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            insert_hr(p, 'C68A6B')
        else:
            # Beautiful Copper Headings
            run.font.size = Pt(13)
            run.font.color.rgb = BRAND_COPPER
            insert_hr(p, 'C68A6B')
            
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        run.font.size = Pt(10.0)
        p.paragraph_format.space_after = Pt(2)

    def embed_image_if_exists(image_source, target_width_inches=6.0, center=True):
        if not image_source:
            return False
            
        if not isinstance(image_source, str):
            image_source = str(image_source)
            
        target_name = image_source.strip()
        if target_name.lower().endswith((".png", ".jpg", ".jpeg")):
            target_name = os.path.splitext(target_name)[0]

        search_dirs = []
        if custom_img_dir:
            search_dirs.append(custom_img_dir)
        search_dirs.append("scraped_images")
            
        for scraped_dir in search_dirs:
            if os.path.exists(scraped_dir):
                for f in os.listdir(scraped_dir):
                    base_f = os.path.splitext(f)[0]
                    if base_f == target_name or f == target_name or f.startswith(target_name + "."):
                        try:
                            img_path = os.path.join(scraped_dir, f)
                            p = doc.add_paragraph()
                            if center:
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            run.add_picture(img_path, width=Inches(target_width_inches))
                            return True
                        except Exception as e:
                            print(f"Failed to embed image {img_path}: {e}")
        return False

    # ==========================================
    # COVER PAGE
    # ==========================================
    doc.add_paragraph() # Spacer
    doc.add_paragraph() # Spacer
    
    brand_p = doc.add_paragraph()
    brand_run = brand_p.add_run("GENERATED BY ARTISANBLUEPRINT")
    brand_run.font.size = Pt(12)
    brand_run.font.color.rgb = BRAND_COPPER
    brand_run.bold = True
    brand_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph() # Spacer

    title_p = doc.add_paragraph()
    title_run = title_p.add_run(project_name)
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = BRAND_DARK
    title_run.bold = True
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    diff_text = plan_data.get("difficulty_level", "DIY Project")
    p_diff = doc.add_paragraph()
    diff_run = p_diff.add_run(f"Difficulty: {diff_text}")
    diff_run.font.size = Pt(14)
    diff_run.font.color.rgb = RGBColor(100, 100, 100)
    p_diff.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph() # Spacer
    if plan_data.get("hero_image_source"):
        embed_image_if_exists(plan_data.get("hero_image_source"), target_width_inches=6.5)

    doc.add_page_break()

    # ==========================================
    # 2. INTRO BOX
    # ==========================================
    if plan_data.get("project_intro"):
        add_heading("Project Overview")
        p_intro = doc.add_paragraph()
        run = p_intro.add_run(plan_data["project_intro"])
        run.italic = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(80, 80, 80)
        doc.add_paragraph() # Spacer

    # ==========================================
    # 3. DIMENSIONS
    # ==========================================
    if plan_data.get("finished_dimensions") or plan_data.get("dimension_image_source"):
        add_heading("Dimensions")
        if plan_data.get("finished_dimensions"):
            p_dim = doc.add_paragraph()
            p_dim.add_run("Finished Dimensions: ").bold = True
            p_dim.add_run(plan_data['finished_dimensions'])

        if plan_data.get("dimension_image_source"):
            embed_image_if_exists(plan_data.get("dimension_image_source"))
        doc.add_paragraph()
    
    # ==========================================
    # 4. SHOPPING LIST (MATERIALS) - TABLE
    # ==========================================
    if plan_data.get("materials"):
        add_heading("Shopping List")
        materials = plan_data.get("materials", [])
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Quantity'
        hdr_cells[1].text = 'Material Description'
        
        # Style Header
        for cell in hdr_cells:
            set_cell_background(cell, "F2E6DF") # Very light copper
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = BRAND_DARK
        
        for material in materials:
            row_cells = table.add_row().cells
            q = material.get("quantity")
            d = material.get("description")
            row_cells[0].text = str(q) if q is not None else "-"
            row_cells[1].text = str(d) if d is not None else ""
            
        doc.add_paragraph()

    # ==========================================
    # 5. CUT LIST - TABLE
    # ==========================================
    if plan_data.get("cut_list"):
        add_heading("Cut List")
        cut_list = plan_data.get("cut_list", [])
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Qty'
        hdr_cells[1].text = 'Dimensions'
        hdr_cells[2].text = 'Part Description'
        
        # Style Header
        for cell in hdr_cells:
            set_cell_background(cell, "F2E6DF")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = BRAND_DARK
        
        for cut in cut_list:
            row_cells = table.add_row().cells
            q = cut.get("quantity")
            dim = cut.get("dimensions")
            d = cut.get("description")
            row_cells[0].text = str(q) if q is not None else "-"
            row_cells[1].text = str(dim) if dim is not None else ""
            row_cells[2].text = str(d) if d is not None else ""
            
        doc.add_paragraph()

    # ==========================================
    # 6. TOOLS
    # ==========================================
    if plan_data.get("tools"):
        add_heading("Tools Required")
        
        tools_page = plan_data.get("tools_image_source")
        if tools_page:
             embed_image_if_exists(tools_page)

        for tool in plan_data.get("tools", []):
            tool_name = tool.get('name') if isinstance(tool, dict) else str(tool)
            if tool_name:
                add_bullet(str(tool_name))
        doc.add_paragraph()
    
    # ==========================================
    # 7. CONSTRUCTION STEPS
    # ==========================================
    def is_heading_line(line_text):
        clean_lower = line_text.lower().strip()
        heading_keywords = [
            "material list", "shopping list", "cutting list", "cut list", "overview",
            "floor framing", "floor deck", "wall framing", "side wall frame", "front/back wall frame",
            "right/left wall frame", "front top wall frame", "rafters", "roof", "siding", "trim",
            "door", "window", "skids", "walls", "floor", "roof deck", "corner trim", "purlins", "blocking"
        ]
        if any(clean_lower == kw or clean_lower.startswith(kw + ":") for kw in heading_keywords):
            return True
        if line_text.isupper() and len(line_text) < 30 and not re.match(r'^\d+', line_text):
            return True
        return False

    def render_custom_line(line_clean):
        if not line_clean:
            return
        if is_heading_line(line_clean):
            add_heading(line_clean, level=2)
        elif line_clean.startswith("•"):
            add_bullet(line_clean.lstrip("•").strip())
        elif re.match(r'^\d+[\s\–\-]+', line_clean) and any(kw in line_clean.lower() for kw in ["2x", "4x", "1x", "plywood", "sheet", "screw", "nail", "shingle", "felt"]):
            add_bullet(line_clean)
        else:
            p = doc.add_paragraph()
            run = p.add_run(line_clean)
            run.font.size = Pt(10.0)
            p.paragraph_format.space_after = Pt(3)

    # ==========================================
    # 7. SPATIAL PAGE-BY-PAGE RENDERING
    # ==========================================
    if plan_data.get("steps"):
        for step_idx, step in enumerate(plan_data.get("steps", [])):
            doc.add_page_break()
            
            # 1. Render Text Lines ABOVE the Image
            lines_above = step.get("lines_above", [])
            for line in lines_above:
                render_custom_line(line)
                
            # 2. Embed Page Diagram Image
            img_candidates = step.get("image_sources", [])
            for candidate in img_candidates:
                embed_image_if_exists(candidate, target_width_inches=6.0)
                
            # 3. Render Text Lines BELOW the Image
            lines_below = step.get("lines_below", [])
            exact_desc = step.get("exact_description")
            
            if lines_below:
                for line in lines_below:
                    render_custom_line(line)
            elif exact_desc:
                for line in str(exact_desc).split("\n"):
                    render_custom_line(line.strip())

    # ==========================================
    # 8. FINISHING INSTRUCTIONS
    # ==========================================
    if plan_data.get("finishing_instructions"):
        doc.add_page_break()
        add_heading("Finishing Instructions", level=1)
        for inst in plan_data.get("finishing_instructions", []):
            if inst:
                add_bullet(str(inst))

    doc.save(output_filename)
    print(f"Generated {output_filename}")

